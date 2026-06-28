from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO


def export_sgpa_report(subjects, sgpa, scale, semester_name="Semester"):
    doc = Document()

    title = doc.add_heading(f"SemSnap — {semester_name} Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

    doc.add_heading("Subject Details", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Subject"
    hdr[1].text = "Credits"
    hdr[2].text = "Grade Point"

    for s in subjects:
        row = table.add_row().cells
        row[0].text = s['name']
        row[1].text = str(s['credits'])
        row[2].text = str(s['grade_point'])

    doc.add_paragraph("")
    doc.add_heading(f"SGPA ({scale} point scale): {sgpa}", level=1)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_cgpa_report(semesters, cgpa, scale):
    doc = Document()

    title = doc.add_heading("SemSnap — CGPA Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

    doc.add_heading("Semester-wise SGPA", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Semester"
    hdr[1].text = "SGPA"
    hdr[2].text = "Credits"

    for i, s in enumerate(semesters, 1):
        row = table.add_row().cells
        row[0].text = f"Semester {i}"
        row[1].text = str(s['sgpa'])
        row[2].text = str(s['credits'])

    doc.add_paragraph("")
    doc.add_heading(f"CGPA ({scale} point scale): {cgpa}", level=1)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
